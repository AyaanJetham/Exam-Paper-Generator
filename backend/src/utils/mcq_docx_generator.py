from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import math


def sanitize(txt):
    if not txt:
        return ""
    return str(txt).replace('\u2019', "'").replace('\u2018', "'") \
                   .replace('\u201c', '"').replace('\u201d', '"') \
                   .replace('\u2013', '-').replace('\u2014', '-')


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run_with_style(para, text, bold=False, size=11):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return run


def generate_mcq_docx(data, output_path):
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_style(title_para, '[ UNIVERSITY EXAMINATION ]', bold=True, size=16)

    subject_name = sanitize(data.get('subject_name', ''))
    if subject_name:
        subj_para = doc.add_paragraph()
        subj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(subj_para, subject_name, bold=True, size=14)

    doc.add_paragraph()

    # Duration/Marks table
    dur_table = doc.add_table(rows=1, cols=2)
    dur_table.style = 'Table Grid'
    for cell in dur_table.rows[0].cells:
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        cell._tc.get_or_add_tcPr().append(tcBorders)
    dur_table.rows[0].cells[0].text = 'Duration: 3 Hours'
    dur_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    dur_table.rows[0].cells[1].text = 'Marks: 80 Marks'
    p = dur_table.rows[0].cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    add_horizontal_rule(doc)
    doc.add_paragraph()

    # Instructions box
    nb_table = doc.add_table(rows=1, cols=1)
    nb_table.style = 'Table Grid'
    nb_cell = nb_table.rows[0].cells[0]

    nb_p = nb_cell.add_paragraph()
    add_run_with_style(nb_p, 'N.B.:', bold=True, size=11)

    for inst in [
        "(1) All questions are compulsory.",
        "(2) Each question carries 2 marks.",
        "(3) Choose the most appropriate answer from the given options.",
        "(4) Total Marks: 80  |  Total Questions: 40",
    ]:
        ip = nb_cell.add_paragraph()
        ip.paragraph_format.left_indent = Cm(0.5)
        add_run_with_style(ip, sanitize(inst), size=10)

    nb_cell.paragraphs[0]._element.getparent().remove(nb_cell.paragraphs[0]._element)

    doc.add_paragraph()

    # MCQ Questions
    mcqs = data.get('mcqs', [])
    for mcq in mcqs:
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(6)
        add_run_with_style(q_para, f"Q.{mcq['qst_num']}.  ", bold=True, size=11)
        add_run_with_style(q_para, sanitize(mcq.get('text', '')), size=11)
        q_para.add_run('\t')
        mr = q_para.add_run('2 Marks')
        mr.font.size = Pt(10)
        q_para.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), WD_ALIGN_PARAGRAPH.RIGHT)

        options = mcq.get('options', {})
        for label in ['A', 'B', 'C', 'D']:
            opt_text = sanitize(options.get(label, ''))
            if opt_text:
                opt_para = doc.add_paragraph()
                opt_para.paragraph_format.left_indent = Cm(1.0)
                add_run_with_style(opt_para, f'{label}. {opt_text}', size=10)

    # ===== ANSWER KEY =====
    doc.add_page_break()

    ak_title = doc.add_paragraph()
    ak_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_style(ak_title, 'ANSWER KEY', bold=True, size=14)
    doc.add_paragraph()

    cols = 4
    per_col = math.ceil(len(mcqs) / cols)

    # Build answer key table
    ak_table = doc.add_table(rows=per_col + 1, cols=cols * 2)
    ak_table.style = 'Table Grid'

    # Header
    headers = ['Q.No', 'Ans'] * cols
    for i, h in enumerate(headers):
        cell = ak_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Answers
    for row_idx in range(per_col):
        for col_idx in range(cols):
            mcq_idx = col_idx * per_col + row_idx
            q_cell = ak_table.rows[row_idx + 1].cells[col_idx * 2]
            a_cell = ak_table.rows[row_idx + 1].cells[col_idx * 2 + 1]
            if mcq_idx < len(mcqs):
                mcq = mcqs[mcq_idx]
                q_cell.text = f"Q.{mcq['qst_num']}"
                a_cell.text = str(mcq.get('answer', '-'))
            q_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            a_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path
