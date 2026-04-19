from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_rule(doc):
    """Add a horizontal line to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run_with_style(para, text, bold=False, size=11, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def sanitize(txt):
    if not txt:
        return ""
    return str(txt).replace('\u2019', "'").replace('\u2018', "'") \
                   .replace('\u201c', '"').replace('\u201d', '"') \
                   .replace('\u2013', '-').replace('\u2014', '-')


def generate_question_paper_docx(data, output_path, difficulty='medium', include_answers=False):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # === HEADER: University Examination ===
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_style(title_para, '[ UNIVERSITY EXAMINATION ]', bold=True, size=16)

    # Subject name
    subject_name = sanitize(data.get('subject_name', ''))
    if subject_name:
        subj_para = doc.add_paragraph()
        subj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(subj_para, subject_name, bold=True, size=14)

    doc.add_paragraph()  # spacer

    # Duration | Marks — two-column borderless table
    dur_table = doc.add_table(rows=1, cols=2)
    dur_table.style = 'Table Grid'
    for cell in dur_table.rows[0].cells:
        cell._tc.get_or_add_tcPr()
        # Remove borders
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        cell._tc.tcPr.append(tcBorders)

    dur_table.rows[0].cells[0].text = 'Duration: 3 Hours'
    dur_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    dur_table.rows[0].cells[1].text = 'Marks: 80 Marks'
    p = dur_table.rows[0].cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].bold = True

    add_horizontal_rule(doc)
    doc.add_paragraph()

    # === INSTRUCTIONS BOX ===
    nb_table = doc.add_table(rows=1, cols=1)
    nb_table.style = 'Table Grid'
    nb_cell = nb_table.rows[0].cells[0]

    nb_para = nb_cell.add_paragraph()
    add_run_with_style(nb_para, 'N.B.:', bold=True, size=11)

    instructions = [
        "(1) Question No 1 is Compulsory.",
        "(2) Attempt any three questions out of the remaining five.",
        "(3) All questions carry equal marks.",
        "(4) Assume suitable data, if required and state it clearly."
    ]
    for inst in instructions:
        inst_para = nb_cell.add_paragraph()
        inst_para.paragraph_format.left_indent = Cm(0.5)
        add_run_with_style(inst_para, sanitize(inst), size=10)

    # Remove the first empty paragraph in cell
    nb_cell.paragraphs[0]._element.getparent().remove(nb_cell.paragraphs[0]._element)

    doc.add_paragraph()  # spacer

    # === QUESTIONS ===
    for q in data.get('paper', []):
        qnum = q.get('qst_num')
        instruction = sanitize(q.get('instruction', ''))

        # Q.X row
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(8)

        # Determine marks display
        marks_display = "20"
        if difficulty == 'balanced':
            if qnum == 1:
                marks_display = "20 (2x10)" 
            elif qnum == 2:
                marks_display = "20 (4x5)"
            elif qnum == 6:
                marks_display = "20 (2x10)"
        else:
            if qnum == 1:
                marks_display = "20 (4x5)"
            elif qnum == 6:
                marks_display = "20 (2x10)"

        add_run_with_style(q_para, f'Q.{qnum}  ', bold=True, size=12)
        if instruction:
            add_run_with_style(q_para, instruction, size=12)

        # Marks — right-aligned via tab
        q_para.add_run('\t')
        marks_run = q_para.add_run(marks_display)
        marks_run.bold = True
        marks_run.font.size = Pt(12)
        q_para.paragraph_format.tab_stops.add_tab_stop(Inches(5.5), WD_ALIGN_PARAGRAPH.RIGHT)

        # Parts
        for part in q.get('parts', []):
            part_text = sanitize(part.get('text', '').strip())
            part_label = part.get('part_label', '')
            part_marks = part.get('marks', '')

            # Defensive formatting: Q1 should be 5x5 even if marks omitted by the model
            if (not part_marks) and qnum == 1 and "4x5" in str(marks_display) and part_text:
                part_marks = "5"

            if part_text:
                p_para = doc.add_paragraph()
                p_para.paragraph_format.left_indent = Cm(1.0)
                add_run_with_style(p_para, f'{part_label}. {part_text}', size=11)
                if part_marks:
                    p_para.add_run('\t')
                    mr = p_para.add_run(str(part_marks))
                    mr.font.size = Pt(11)
                    p_para.paragraph_format.tab_stops.add_tab_stop(Inches(5.0), WD_ALIGN_PARAGRAPH.RIGHT)

            # Subparts
            for idx, subpart in enumerate(part.get('subparts', [])):
                sub_text = sanitize(subpart.get('text', ''))
                sub_label = subpart.get('sub_label', '')
                sub_marks = subpart.get('marks', '')

                sp_para = doc.add_paragraph()
                sp_para.paragraph_format.left_indent = Cm(1.5)

                if not part_text and idx == 0:
                    label_str = f'{part_label}.      {sub_label}. {sub_text}'
                else:
                    label_str = f'{sub_label}. {sub_text}'

                add_run_with_style(sp_para, label_str, size=10)
                if sub_marks:
                    sp_para.add_run('\t')
                    mr = sp_para.add_run(str(sub_marks))
                    mr.font.size = Pt(10)
                    sp_para.paragraph_format.tab_stops.add_tab_stop(Inches(5.0), WD_ALIGN_PARAGRAPH.RIGHT)

    # ===== OPTIONAL ANSWER KEY SECTION =====
    if include_answers:
        doc.add_page_break()
        ans_title = doc.add_paragraph()
        ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(ans_title, 'MODEL ANSWER KEY / SOLUTION GUIDE', bold=True, size=16)
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = add_run_with_style(info_para, '(Includes Marking Scheme and Technical Key Points)', size=10)
        run.italic = True
        
        add_horizontal_rule(doc)
        doc.add_paragraph()

        for q in data.get('paper', []):
            qnum = q.get('qst_num')
            head_p = doc.add_paragraph()
            add_run_with_style(head_p, f"Question {qnum}:", bold=True, size=12)
            
            for part in q.get('parts', []):
                part_label = part.get('part_label', '')
                ans = sanitize(part.get('answer_guide', '').strip())
                
                subparts = part.get('subparts', [])
                if subparts:
                    for sub in subparts:
                        sub_label = sub.get('sub_label', '')
                        sub_ans = sanitize(sub.get('answer_guide', '').strip())
                        if sub_ans:
                            ans_p = doc.add_paragraph()
                            ans_p.paragraph_format.left_indent = Cm(1.0)
                            add_run_with_style(ans_p, f"({part_label}.{sub_label}) ", bold=True, size=11)
                            add_run_with_style(ans_p, sub_ans, size=11)
                elif ans:
                    ans_p = doc.add_paragraph()
                    ans_p.paragraph_format.left_indent = Cm(1.0)
                    add_run_with_style(ans_p, f"({part_label}) ", bold=True, size=11)
                    add_run_with_style(ans_p, ans, size=11)

            doc.add_paragraph()

    doc.save(output_path)
    return output_path
